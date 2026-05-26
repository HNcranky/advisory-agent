"""One-shot inspection helper for HUST pre-flight (Plan 01).

Subcommands:
  listing      Inspect program listing for published quotas.
  proposal     Crawl ts.hust.edu.vn/b/de-an-tuyen-sinh for 2026 PDFs.
  article URL  Inspect a single article URL for embedded PDFs.
  pdf URL      Run pdfplumber on a PDF and dump first-page samples.
  brochure     Resolve PDF from nxbbachkhoa.vn ebook viewer.
"""
from __future__ import annotations

import io
import re
import sys
from collections import Counter

from bs4 import BeautifulSoup

from ingestion.fetchers.http_fetcher import http_fetch
from ingestion.parsers.hust_program_parser import HustProgramParser

LISTING_URL = "https://ts.hust.edu.vn/training-cate/nganh-dao-tao-dai-hoc"
PROPOSAL_LIST_URL = "https://ts.hust.edu.vn/b/de-an-tuyen-sinh"
BROCHURE_URL = "https://nxbbachkhoa.vn/ebook-free/12397/0/1"


def cmd_listing() -> None:
    r = http_fetch(LISTING_URL)
    print(f"status={r.http_status} size={len(r.raw_content)}")

    parser = HustProgramParser()
    facts = parser.parse(r.raw_content, source_url=r.url)
    print(f"facts produced by parser: {len(facts)}")

    html = r.raw_content.decode("utf-8", errors="replace")
    soup = BeautifulSoup(html, "html.parser")

    pat = re.compile(r"Ch[ỉi]\s*ti[êe]u tuy[ểe]n sinh:")
    quota_buckets: Counter[str] = Counter()
    rows: list[tuple[str, str]] = []

    for el in soup.find_all(string=pat):
        parent = el.parent
        strong = None
        if parent is not None:
            strong = parent.find("strong")
            if strong is None:
                strong = parent.find_next("strong")
        qtxt = strong.get_text(strip=True) if strong is not None else ""
        quota_buckets[qtxt] += 1

        title = ""
        node = parent
        for _ in range(8):
            if node is None:
                break
            heading = node.find_previous(["h2", "h3", "h4"])
            if heading is not None:
                title = heading.get_text(strip=True)
                break
            node = node.parent
        rows.append((title, qtxt))

    print("\n-- quota frequency --")
    for q, n in quota_buckets.most_common():
        print(f"  {q!r}: {n}")

    print("\n-- programs WITH a non-empty quota --")
    for title, qtxt in rows:
        if qtxt and qtxt != "0":
            print(f"  quota={qtxt!r:>5}  program={title[:140]!r}")

    print("\n-- sample of programs with EMPTY quota (first 5) --")
    empties = [r for r in rows if not r[1]]
    for title, qtxt in empties[:5]:
        print(f"  quota={qtxt!r:>5}  program={title[:140]!r}")
    print(f"  (... {len(empties)} programs with empty quota total)")


def cmd_proposal() -> None:
    r = http_fetch(PROPOSAL_LIST_URL)
    print(f"status={r.http_status} size={len(r.raw_content)}")
    soup = BeautifulSoup(r.raw_content, "html.parser")
    print("Title:", soup.title.get_text() if soup.title else "?")
    print()

    pdfs = sorted({a["href"] for a in soup.find_all("a", href=True) if a["href"].lower().endswith(".pdf")})
    print(f"-- PDF links on the listing ({len(pdfs)}) --")
    for p in pdfs:
        print(" ", p)

    print()
    print("-- anchors mentioning 2026 --")
    seen = set()
    for a in soup.find_all("a", href=True):
        txt = a.get_text(strip=True)
        if "2026" in txt:
            key = (a["href"], txt[:120])
            if key in seen:
                continue
            seen.add(key)
            print(f"  {a['href']}")
            print(f"    text: {txt[:160]}")

    # Also dump all article headings (de-an articles are usually <h2>/<h3>).
    print()
    print("-- all article-like anchors (first 30) --")
    n = 0
    for a in soup.find_all("a", href=True):
        href = a["href"]
        if "/b/" in href and href not in ("/b/de-an-tuyen-sinh",):
            txt = a.get_text(strip=True)
            if not txt:
                continue
            n += 1
            if n <= 30:
                print(f"  {href}  | {txt[:120]}")
    print(f"  (total /b/ anchors: {n})")


def cmd_article(url: str) -> None:
    r = http_fetch(url)
    print(f"status={r.http_status} size={len(r.raw_content)} url={r.url}")
    soup = BeautifulSoup(r.raw_content, "html.parser")
    print("Title:", soup.title.get_text() if soup.title else "?")

    print("\n-- PDF/document links inside the article --")
    seen = set()
    for a in soup.find_all("a", href=True):
        href = a["href"]
        low = href.lower()
        if low.endswith(".pdf") or "pdf" in low or low.endswith(".docx") or low.endswith(".doc"):
            if href in seen:
                continue
            seen.add(href)
            print(f"  {href}")
            print(f"    anchor: {a.get_text(strip=True)[:160]}")

    print("\n-- embeds/iframes --")
    for tag in soup.find_all(["embed", "iframe", "object"]):
        src = tag.get("src") or tag.get("data")
        if src:
            print(f"  {tag.name}: {src}")


def cmd_pdf(url: str) -> None:
    import pdfplumber

    r = http_fetch(url)
    print(f"status={r.http_status} size={len(r.raw_content)} url={r.url}")
    if r.http_status != 200:
        return
    with pdfplumber.open(io.BytesIO(r.raw_content)) as pdf:
        print(f"pages: {len(pdf.pages)}")
        max_pages = min(len(pdf.pages), 8)
        for i in range(max_pages):
            text = pdf.pages[i].extract_text() or ""
            print(f"--- page {i+1} ({len(text)} chars) ---")
            print(text[:500])
            print()


def cmd_html2026(url: str) -> None:
    r = http_fetch(url)
    print(f"status={r.http_status} size={len(r.raw_content)} url={r.url}")
    soup = BeautifulSoup(r.raw_content, "html.parser")
    main = soup.find("article") or soup.find("main") or soup.body
    text = main.get_text(separator="\n", strip=True) if main else ""
    lines = text.split("\n")
    print(f"total lines: {len(lines)}")

    keyword = re.compile(r"ch[ỉi].{0,3}ti[êe]u|tuy[ểe]n sinh|chỉ tiêu", re.I)
    hits = [l for l in lines if keyword.search(l)]
    print(f"lines mentioning chi tieu/tuyen sinh: {len(hits)}")
    print()
    print("-- first 30 chi-tieu/tuyen-sinh lines --")
    for l in hits[:30]:
        print(repr(l[:240]))

    tables = soup.find_all("table")
    print()
    print(f"-- tables on page: {len(tables)} --")
    for i, t in enumerate(tables[:5]):
        rows = t.find_all("tr")
        print(f"--- table {i+1} ({len(rows)} rows) ---")
        for row in rows[:8]:
            cells = [c.get_text(strip=True)[:80] for c in row.find_all(["td", "th"])]
            print("  | ".join(cells))


def cmd_html2026_full(url: str) -> None:
    r = http_fetch(url)
    soup = BeautifulSoup(r.raw_content, "html.parser")
    tables = soup.find_all("table")
    for i, t in enumerate(tables):
        rows = t.find_all("tr")
        print(f"=== table {i+1} ({len(rows)} rows) ===")
        for j, row in enumerate(rows):
            cells = [c.get_text(strip=True) for c in row.find_all(["td", "th"])]
            print(f"  [{j:>2}] " + " | ".join(cells))
        print()


def cmd_docx(url: str) -> None:
    import docx

    r = http_fetch(url)
    print(f"status={r.http_status} size={len(r.raw_content)} url={r.url}")
    if r.http_status != 200:
        return
    d = docx.Document(io.BytesIO(r.raw_content))
    print(f"paragraphs: {len(d.paragraphs)}")
    print(f"tables: {len(d.tables)}")
    print()
    print("-- first 40 non-empty paragraphs --")
    n = 0
    for p in d.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        n += 1
        if n <= 40:
            print(f"  {txt[:240]}")
    print(f"  (total non-empty paragraphs: {n})")
    print()
    print("-- tables --")
    for i, t in enumerate(d.tables[:10]):
        rows = t.rows
        print(f"--- table {i+1} ({len(rows)} rows x {len(rows[0].cells) if rows else 0} cols) ---")
        for row in rows[:10]:
            cells = [c.text.strip()[:80] for c in row.cells]
            print("  | ".join(cells))
        if len(rows) > 10:
            print(f"  (... {len(rows)-10} more rows)")


def cmd_brochure() -> None:
    r = http_fetch(BROCHURE_URL)
    print(f"status={r.http_status} size={len(r.raw_content)} url={r.url}")
    html = r.raw_content.decode("utf-8", errors="replace")
    candidates = sorted(set(re.findall(r'https?://[^\s"\'<>]+\.pdf', html, re.I)))
    print(f"\n-- inline PDF URLs ({len(candidates)}) --")
    for c in candidates:
        print(" ", c)
    soup = BeautifulSoup(r.raw_content, "html.parser")
    print("\n-- embed/iframe/object srcs --")
    for tag in soup.find_all(["embed", "iframe", "object"]):
        src = tag.get("src") or tag.get("data")
        if src:
            print(f"  {tag.name}: {src}")
    # Look for typical ebook viewer hints
    print("\n-- script tags containing 'pdf' or 'file' (first 5 matches) --")
    n = 0
    for s in soup.find_all("script"):
        txt = (s.string or s.get_text() or "")
        if "pdf" in txt.lower() or "file" in txt.lower():
            n += 1
            if n <= 5:
                snippet = txt[:600]
                print(f"  ---script {n}---")
                print(f"  {snippet}")


def main() -> None:
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(2)
    cmd, *rest = sys.argv[1:]
    if cmd == "listing":
        cmd_listing()
    elif cmd == "proposal":
        cmd_proposal()
    elif cmd == "article":
        cmd_article(rest[0])
    elif cmd == "pdf":
        cmd_pdf(rest[0])
    elif cmd == "brochure":
        cmd_brochure()
    elif cmd == "html2026":
        cmd_html2026(rest[0])
    elif cmd == "docx":
        cmd_docx(rest[0])
    elif cmd == "html2026full":
        cmd_html2026_full(rest[0])
    else:
        print(f"unknown command: {cmd}")
        print(__doc__)
        sys.exit(2)


if __name__ == "__main__":
    main()
